from docx import Document
from openai import OpenAI
from transcribe import Transcript
import Config 

Trans = Transcript()
file = Transcript.transcription()
# file = open('Transcript.txt','r')

client = OpenAI(
  organization=Config.organization,
  project=Config.project,
  api_key=Config.gpt_api_key
)

completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": "You are a helpful assistant."},
        {
            "role": "user",  
            "content": 
            f"""
I have gone and found the top 10 most popular youtube videos under the EA FC game genre in the past week.

I will now give you the transcripts of the top 10 videos, and after analyzing the transcripts in order
to understand the context of the video I want you to provide the following: 

[Please use the format given below and create one brief for each transcript]

TEMPLATE:

1. Intro 

- Greet Viewers: Start with an enthusiastic, welcoming tone (e.g., “Hey everyone, welcome back!”). 
- Describe the Video: Briefly explain what the video covers (e.g., “Today, we’re diving into a challenge…”). 
- Highlight Any Stakes or Goals: Mention any specific stakes or objectives (e.g., “If I don’t complete this challenge, I’ll have to…”). 
- First Call-to-Action (CTA): Prompt viewers to like or subscribe (e.g., “If you’re excited, don’t forget to hit like and subscribe!”). 

2. Main Content 

- Outline Content in Parts: Break the main video content into clearly defined parts: 
 - Part Introduction: State what will happen in each part (e.g., “First round - let’s get started!”). 
 - Actions and Commentary: Follow through with actions, adding commentary and reactions as they happen. 
- Transitions Between Parts: Conclude each section with a lead into the next (e.g., “Alright, moving on to the next part…”). 

3. Outro 

- Summarize the Video: Recap main points (e.g., “So, that’s it for today’s challenge!”). 
- Thank Viewers: Show appreciation for viewers’ time (e.g., “Thanks for spending time with me today!”). 
- Final CTA: “Don’t forget to like, comment, and subscribe for more!” 

4. Technical and Style Requirements 

- Quality Check: Ensure clear audio and video quality. Conduct a test recording and share it for approval. 
- Format: Record in horizontal orientation for YouTube. 
- Required Equipment: Webcam for face shots, screen capture for gameplay. 
- Visual Elements: Keep all main action clearly visible with good lighting. 
- Sound: Minimize background noise. 

5. Video Length 

- Expected Duration: The final video should be around 12-15 minutes. 

6. Additional Details 

- Contact for Questions: If any questions arise during production, reach out to the Producer at [Producer's Number]. 
- Deadlines: 
 - Script Draft: [Date] 
 - Footage Recording: [Date] 
 - First Edit: [Date] 
 - Final Cut: [Date] 

- Point of Contact (POC): 
 - Video Editor: [Name/Contact] 
 - Social Media Manager: [Name/Contact] 

 Prompt for AI to Generate a Brief
 “Generate a step-by-step video brief that breaks down content creation into simple tasks. Ensure instructions are clear enough for someone with minimal experience to follow. Divide each video section (intro, main parts, outro) with concise instructions on what to say and do, especially for transitions. Add any CTA lines directly. Include specific technical requirements, such as equipment and video length, so the creator knows exactly what’s needed. Be thorough yet simple, so even a beginner can understand and complete the task.”


EXAMPLE: 

Intro - Greet Viewers: "Hey everyone, welcome back! Today, we’ve got a really fun challenge ahead!"
 - Describe the Video: "I’m going to try to guess players based on their FIFA card histories to build my ultimate FC25 team." 
- Highlight Any Stakes or Goals: "Here's the twist—if I don’t guess correctly, I’ll be limited to a lower-rated player. And, if I lose a match at the end with my new team, I’ll have to discard one player for every goal I lose by." 
- First Call-to-Action (CTA): "If you’re ready for some high-stakes team-building, hit that like button and subscribe for more!" 

 Main Content Outline 

Part 1 - Building My Attack 
 - Part Introduction: "Alright, let’s start with my first striker. I’ll be guessing the player based on their FIFA card ratings from past years. Let’s see if I can get this right in one try!" - Actions and Commentary: Proceed with the guess (e.g., "This player has gotten better every year since FIFA 14… I think it’s Bernardo Silva!"). After the guess, add a reaction: "Yes! I got it! Adding a top-tier striker to the team." 
 - Transition: "Alright, that’s our first striker sorted. Next, let’s get a CAM on board." 

Part 2 - CAM Challenge 
 - Part Introduction: "Now, I’m playing for a CAM position. The stakes are the same—guess correctly and I get a top-rated player; get it wrong, and it’s a downgrade." 
 - Actions and Commentary: Narrate your thought process as you guess, using lifelines if necessary. React to each decision and lifeline use. 
 - Transition: "Great! Our CAM is set. Let’s move on to the next part." 

Part 3 - The Lifeline Dilemma 
 - Part Introduction: "Alright, now things get a little tougher. This time, I’m playing for a left-back, and I only have a few lifelines left. I’ll need to be strategic with my guesses." 
 - Actions and Commentary: Narrate as you narrow down your choices, discussing the use of lifelines to keep viewers engaged.
 - Transition: "Managed to get my left-back! Let’s keep the momentum going and move on to the center-back."
 
Part 4 - Viewers’ Challenge
 - Part Introduction: "This one’s for you! Let me know in the comments who you think this player is based on their FIFA history." 
 - Actions and Commentary: Provide hints for viewers and encourage them to guess in the comments.
 - Transition: "While you’re figuring that out, I’ll move on to get our final center-back." 

5. Final Gameplay 

 - Part Introduction: "Alright, time for the ultimate test! Now that I’ve got my team, I’ll play one match to see how we stack up." 
 - Actions and Commentary: Provide lively commentary during the gameplay, showing the stakes of discarding players based on goals conceded. 
 - Transition: "Let’s see if I can hold my own with this team!"

 Outro 

- Summarize the Video: "Alright, that wraps up the team-building challenge for today. I took some big risks, used up a few lifelines, but ended up with a pretty solid squad." 
- Thank Viewers: "Thanks for joining me on this journey. Your support keeps these challenges exciting!" 
- Final CTA: "If you enjoyed this, don’t forget to like, comment, and subscribe for more content!" 

 Technical and Style Requirements 

- Quality Check: Ensure clear audio and video quality.
 - Format: Record in horizontal orientation for YouTube. 
- Required Equipment: Use a webcam for face shots and screen capture for gameplay. 
- Visual Elements: Keep the main action clearly visible with good lighting. 
- Sound: Minimize background noise; light background music if preferred. Video Length 
- Expected Duration: Aim for a final length of 12-15 minutes. Additional Details
 - Contact for Questions: Reach out to the Producer at [Producer's Number]. 
- Deadlines: - Script Draft: [Date] - Footage Recording: [Date] - First Edit: [Date] - Final Cut: [Date]

 - Point of Contact (POC):
 - Video Editor: [Name/Contact]
 - Social Media Manager: [Name/Contact]

below are the transcripts of the videos:
{file.read()}
                                        """
        }
    ]
)

reply = completion.choices[0].message.content
try:
    transcript_file = open('transcript.txt','r')
    # Create a Word document and add the reply
    document = Document()
    document.add_paragraph(reply)
    document.add_paragraph("----BELOW ARE THE FULL TRANSCRIPT FOR ALL THE VIDEOS----")
    document.add_paragraph(transcript_file.read())
    document.add_paragraph("Prepared By: Saad Amawi")


    # Save the document
    document.save("ChatGPT_Reply3.docx")

    print("The reply has been saved to ChatGPT_Reply.docx")
except Exception as e:
    print("error adding transcriptions to document",e)

